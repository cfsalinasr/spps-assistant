"""python-docx DOCX generator for SPPS synthesis guides."""

from pathlib import Path
from typing import Dict, List, Optional

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from spps_assistant.domain.constants import THREE_LETTER_CODE
from spps_assistant.domain.models import (
    CouplingCycle, CycleGuideViewData, CyclePageData, DispatchRow, GmpStep, SecondaryCouplingRow,
    SynthesisConfig, Vessel, VesselAssignment, YieldResult, SolubilityResult
)

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------



def _set_cell_bg(cell, hex_color: str) -> None:
    """Set a table cell background colour."""
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tc_pr.append(shd)


def _set_cell_borders(table) -> None:
    """Add thin borders to all cells in a table."""
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_borders = OxmlElement('w:tcBorders')
            for border_name in ('top', 'left', 'bottom', 'right'):
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '888888')
                tc_borders.append(border)
            tc_pr.append(tc_borders)


def _style_header_row(row, bg_hex: str = '2C3E50') -> None:
    """Style a table header row (dark background, white bold text)."""
    for cell in row.cells:
        _set_cell_bg(cell, bg_hex)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9)


def _fill_table_row(row, row_data: List[str], n_cols: int) -> None:
    """Fill a single table row with text content."""
    for c_idx, cell_text in enumerate(row_data):
        if c_idx < n_cols:
            cell = row.cells[c_idx]
            para = cell.paragraphs[0]
            run = para.add_run(str(cell_text))
            run.font.size = Pt(9)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _set_col_widths(table, col_widths: List[float]) -> None:
    """Apply column widths (in cm) to a table."""
    for i, width in enumerate(col_widths):
        if i < len(table.columns):
            for cell in table.columns[i].cells:
                cell.width = Cm(width)


def _add_table_with_header(doc: Document, data: List[List[str]],
                            col_widths: Optional[List[float]] = None,
                            header_bg: str = '2C3E50') -> None:
    """Add a table to document with styled header row.

    Args:
        doc: Word Document object
        data: 2D list of strings, first row is header
        col_widths: Optional list of column widths in cm
        header_bg: Header background hex color
    """
    if not data:
        return

    n_rows = len(data)
    n_cols = max(len(row) for row in data)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    if col_widths:
        _set_col_widths(table, col_widths)

    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx]
        _fill_table_row(row, row_data, n_cols)

        if r_idx == 0:
            _style_header_row(row, header_bg)
        elif r_idx % 2 == 0:
            for cell in row.cells:
                _set_cell_bg(cell, 'F8F9FA')

    _set_cell_borders(table)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add a bold heading paragraph to *doc* at the given heading level."""
    para = doc.add_paragraph(text)
    run = para.runs[0] if para.runs else para.add_run(text)
    run.bold = True
    run.font.size = Pt(14 - level * 2)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def _build_cover(doc: Document, synthesis_name: str, date_str: str,
                 vessels: List[Vessel], yield_results: List[YieldResult]) -> None:
    """Append the cover page (title, metadata table, vessel summary) to *doc*."""
    doc.add_heading('SPPS Synthesis Guide', 0)
    doc.add_heading(synthesis_name, 1)

    # Metadata
    meta = doc.add_table(rows=6, cols=2)
    meta.style = 'Table Grid'
    meta_data = [
        ('Prepared:', date_str),
        ('Date signed:', '______________________________'),
        ('Operator:', '______________________________'),
        ('Synthesizer:', '______________________________'),
        ('Project:', '______________________________'),
        ('Page:', '1'),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta.rows[i].cells[0].text = k
        meta.rows[i].cells[1].text = v
        meta.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Vessel summary
    doc.add_heading(f'Synthesis Summary — {len(vessels)} vessel(s)', 2)
    yield_map = {yr.vessel_number: yr for yr in yield_results}

    summary_data = [['#', 'Name', 'Sequence (N→C)', 'Len', 'MW (Da)', 'Yield (mg)']]
    for v in vessels:
        yr = yield_map.get(v.number)
        seq = ''.join(v.original_tokens)
        summary_data.append([
            str(v.number), v.name,
            seq[:40] + ('...' if len(seq) > 40 else ''),
            str(v.length),
            f"{yr.peptide_mw:.1f}" if yr else '—',
            f"{yr.theoretical_yield_mg:.1f}" if yr else '—',
        ])

    _add_table_with_header(doc, summary_data)
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Cycle pages
# ---------------------------------------------------------------------------

def _add_aa_dispatch_table(doc: Document, dispatch_rows: List[DispatchRow]) -> None:
    """Add the AA dispatch table for a cycle page to *doc*."""
    aa_data = [['Residue', 'Fmoc-MW', 'mmol', 'Volume (mL)', 'Formula', 'Status', 'Vessels']]
    for row in dispatch_rows:
        aa_data.append([
            row.residue_3letter, f"{row.fmoc_mw:.1f}", f"{row.mmol:.4f}", f"{row.volume_ml:.3f}",
            row.formula_shown, '[ ]',
            ', '.join(str(vn) for vn in row.vessel_numbers),
        ])
    _add_table_with_header(doc, aa_data)


def _add_deprotection_table(doc: Document, deprotection_steps: List[GmpStep]) -> None:
    """Add the deprotection steps table for a cycle page to *doc*."""
    dep_data = [['[ ]', 'Step', 'Details', 'Time']]
    for step in deprotection_steps:
        dep_data.append(['[ ]', step.label, step.detail, step.duration])
    _add_table_with_header(doc, dep_data, header_bg='1A5276')


def _add_coupling_table(doc: Document, coupling_steps: List[GmpStep]) -> None:
    """Add the coupling steps table for a cycle page to *doc*."""
    coup_data = [['[ ]', 'Step', 'Details', 'Time']]
    for step in coupling_steps:
        checkbox = '[ ]' if step.n_checkboxes > 0 else ''
        coup_data.append([checkbox, step.label, step.detail, step.duration])
    _add_table_with_header(doc, coup_data, header_bg='1E8449')


def _add_vessel_assignment(doc: Document, vessel_assignments: List[VesselAssignment], vessel_label: str) -> None:
    """Add vessel assignment bullet lines to doc."""
    p = doc.add_paragraph()
    p.add_run('Vessel Assignment:').bold = True
    for va in vessel_assignments:
        residue = va.residue_3letter if va.residue_3letter is not None else 'OUT'
        text = f"  {vessel_label} {va.vessel_number} [{va.vessel_name}]: {residue}"
        doc.add_paragraph(text, style='List Bullet')


def _add_secondary_coupling_table(
    doc: Document, secondary_coupling_rows: Optional[List[SecondaryCouplingRow]]
) -> None:
    """Add secondary coupling verification table (Teabag method only)."""
    if secondary_coupling_rows is None:
        return
    doc.add_heading('Secondary Coupling Verification', 3)
    sec_data = [['Vessel #', 'Name', 'Residue', '1st [ ]', '2nd [ ]', '3rd [ ]', '4th [ ]']]
    for row in secondary_coupling_rows:
        sec_data.append([str(row.vessel_number), row.vessel_name, row.residue_3letter, '[ ]', '[ ]', '[ ]', '[ ]'])
    _add_table_with_header(doc, sec_data)


def _add_cycle_page(doc: Document, cycle_page: CyclePageData, vessel_label: str) -> None:
    """Add a single coupling cycle page to the DOCX document."""
    p = doc.add_paragraph()
    run = p.add_run(
        f"Cycle {cycle_page.cycle_number} of {cycle_page.total_cycles}  |  "
        f"Date: ________________  |  Operator: ________________"
    )
    run.bold = True
    run.font.size = Pt(10)

    doc.add_heading(f'Cycle {cycle_page.cycle_number} — AA Dispatch', 3)
    _add_aa_dispatch_table(doc, cycle_page.dispatch_rows)
    doc.add_paragraph()

    doc.add_heading('Deprotection', 3)
    _add_deprotection_table(doc, cycle_page.deprotection_steps)
    doc.add_paragraph()

    doc.add_heading('Coupling', 3)
    _add_coupling_table(doc, cycle_page.coupling_steps)
    doc.add_paragraph()

    _add_vessel_assignment(doc, cycle_page.vessel_assignments, vessel_label)
    _add_secondary_coupling_table(doc, cycle_page.secondary_coupling_rows)


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_cycle_guide_docx(
    path: Path,
    synthesis_name: str,
    date_str: str,
    vessels: List[Vessel],
    cycle_guide_data: CycleGuideViewData,
    config: SynthesisConfig,
    yield_results: List[YieldResult],
) -> None:
    """Generate a GMP-compliant cycle guide DOCX document.

    cycle_guide_data is precomputed and shared with the PDF exporter and the
    GUI preview so all three can never drift apart.
    """
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Cover
    _build_cover(doc, synthesis_name, date_str, vessels, yield_results)

    # Coupling cycle pages
    for i, cycle_page in enumerate(cycle_guide_data.cycles):
        _add_cycle_page(doc, cycle_page, config.vessel_label)
        if i < len(cycle_guide_data.cycles) - 1:
            doc.add_page_break()

    doc.save(str(path))


def _add_orthogonal_warning_docx(doc: Document, orthogonal_tokens: List[str]) -> None:
    """Add a red-highlighted warning paragraph for orthogonal protecting groups."""
    from spps_assistant.domain.constants import ORTHOGONAL_PROTECTING_GROUPS
    from spps_assistant.domain.sequence import parse_token

    p = doc.add_paragraph()
    run = p.add_run('\u26a0  ORTHOGONAL PROTECTING GROUP — ADDITIONAL POST-SYNTHESIS STEP REQUIRED')
    run.bold = True
    run.font.color.rgb = RGBColor(0x7B, 0x24, 0x1C)
    run.font.size = Pt(9)

    for tok in orthogonal_tokens:
        try:
            _, prot = parse_token(tok)
        except ValueError:
            prot = ''
        info = ORTHOGONAL_PROTECTING_GROUPS.get(prot, {})
        display = info.get('display', tok)
        msg = info.get('warning', 'Requires special deprotection — not removed by TFA.')
        bp = doc.add_paragraph(style='List Bullet')
        run2 = bp.add_run(f'{display} ({tok}): ')
        run2.bold = True
        run2.font.color.rgb = RGBColor(0x7B, 0x24, 0x1C)
        run2.font.size = Pt(9)
        run3 = bp.add_run(msg)
        run3.font.color.rgb = RGBColor(0x7B, 0x24, 0x1C)
        run3.font.size = Pt(9)


def generate_peptide_info_docx(
    path: Path,
    synthesis_name: str,
    vessels: List[Vessel],
    solubility_results: Dict[int, SolubilityResult],
    yield_results: List[YieldResult],
) -> None:
    """Generate peptide physicochemical info DOCX."""
    path = Path(path)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading(f'Peptide Information — {synthesis_name}', 0)

    yield_map = {yr.vessel_number: yr for yr in yield_results}

    for vessel in vessels:
        sol = solubility_results.get(vessel.number)
        yr = yield_map.get(vessel.number)

        doc.add_heading(f'Vessel {vessel.number}: {vessel.name}', 1)

        seq_str = ''.join(vessel.original_tokens)
        rev_str = ''.join(vessel.reversed_tokens)

        info_data = [
            ['Property', 'Value'],
            ['Sequence (N→C)', seq_str],
            ['Synthesis order (C→N)', rev_str],
            ['Length', str(vessel.length)],
            ['Peptide MW (Da)', f"{yr.peptide_mw:.2f}" if yr else '—'],
            ['Theoretical Yield (mg)', f"{yr.theoretical_yield_mg:.2f}" if yr else '—'],
            ['Resin mass (g)', f"{vessel.resin_mass_g:.4f}"],
            ['Resin substitution (mmol/g)', f"{vessel.substitution_mmol_g:.4f}"],
            ['Yield formula', yr.formula_shown if yr else '—'],
        ]

        if sol:
            info_data += [
                ['GRAVY score', f"{sol.gravy:.3f}" if sol.gravy is not None else '—'],
                ['Net charge (pH 7)', f"{sol.net_charge_ph7:.2f}" if sol.net_charge_ph7 is not None else '—'],
                ['pI', f"{sol.p_i:.2f}" if sol.p_i is not None else '—'],
                ['Hydrophobicity (KD)', f"{sol.kd_avg:.3f}"],
                ['Hydrophobicity (Eisenberg)', f"{sol.eisenberg_avg:.3f}"],
                ['Hydrophobicity (Black & Mould)', f"{sol.black_mould_avg:.3f}"],
                ['Classification', 'Hydrophobic' if sol.is_hydrophobic else 'Hydrophilic'],
                ['Light sensitive', 'Yes' if sol.light_sensitive else 'No'],
                ['Solubilization', sol.recommendation],
            ]

        _add_table_with_header(doc, info_data)

        if sol and sol.orthogonal_groups:
            _add_orthogonal_warning_docx(doc, sol.orthogonal_groups)

        doc.add_paragraph()

    doc.save(str(path))
